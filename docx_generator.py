from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import Dict, Any, Optional

class DOCXGenerator:
    
    def __init__(self):
        pass
    
    def generate_receipt(self, data: Dict[str, Any]) -> Optional[BytesIO]:
        try:
            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
            
            self._create_header(doc, data)
            self._create_main_content(doc, data)
            self._create_footer(doc, data)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            print(f"Erro ao gerar DOCX: {e}")
            return None
    
    def _create_header(self, doc, data):
        header = doc.add_heading('CAPA DE RECEBIMENTO DE FRETE', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        loja_num = data.get('loja', 'N/A')
        store_info = doc.add_paragraph()
        store_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = store_info.add_run(f'LOJA: {loja_num}')
        run.bold = True
        run.font.size = Inches(0.16)
    
    def _create_main_content(self, doc, data):
        doc.add_paragraph()
        
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('DESTINATÁRIO:', data.get('destinatario_nome', 'N/A')),
            ('ENDEREÇO:', f"{data.get('destinatario_endereco', 'N/A')}, {data.get('destinatario_bairro', 'N/A')}"),
            ('CIDADE/UF:', f"{data.get('destinatario_municipio', 'N/A')}/{data.get('destinatario_uf', 'N/A')}"),
            ('CEP:', data.get('destinatario_cep', 'N/A')),
            ('REMETENTE:', data.get('remetente_nome', 'N/A')),
            ('ORIGEM:', f"{data.get('remetente_municipio', 'N/A')}/{data.get('remetente_uf', 'N/A')}"),
            ('DATA EMISSÃO:', data.get('data_emissao', 'N/A')),
            ('VALOR TOTAL:', f"R$ {data.get('valor_total', '0,00')}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            for cell in row.cells:
                cell.width = Inches(3)
    
    def _create_footer(self, doc, data):
        doc.add_paragraph()
        
        footer_table = doc.add_table(rows=2, cols=2)
        footer_table.style = 'Table Grid'
        
        footer_table.rows[0].cells[0].text = f"NF-e Nº: {data.get('numero_nfe', 'N/A')}"
        footer_table.rows[0].cells[1].text = f"SÉRIE: {data.get('serie', 'N/A')}"
        footer_table.rows[1].cells[0].text = f"VOLUME: {data.get('volume_number', '1/1')}"
        footer_table.rows[1].cells[1].text = f"CHAVE: {data.get('chave_acesso', 'N/A')[:20]}..."
        
        for row in footer_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        signature = doc.add_paragraph('ASSINATURA DO RECEBEDOR: _' + '_' * 50)
        signature.alignment = WD_ALIGN_PARAGRAPH.CENTER